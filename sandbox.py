import docx
doc = docx.Document("output.docx")

# print(doc.paragraphs[0].text)
print(doc.paragraphs[0].runs[0])
# print(doc.paragraphs[0].runs[0].text)

# doc = docx.Document()
# create a document

#text_to_be_added = "Hello! This is my text!"

#   doc.add_paragraph("Hello! This is my text!")
#   paraObject = doc.add_paragraph("I love python.")
# makes new paragraph
#   paraObject.add_run(" I am continuing on this paragraph.")
# adds to existing paragraph

#   doc.save("new.docx")
# name your new document
#   doc.save("output.docx")

#print(doc.paragraphs[0].text)
