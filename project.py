import re
import docx

# string.replace_all(",", "")

def input_to_list(input_text):
    full_list = []
    for messy_word in input_text:
        # pulls out the individual words from the data
        step_1 = re.findall(r'\w+', messy_word) 
        # puts the words into a list
        full_list.append(step_1)
    return(full_list)

def compare_lists(input_text, freq_list, number_of_words):
    output = []
    global words_known
    global words_not_known
    global freq_num
    global clean_input

    freq_num = number_of_words

    for paragraph in input_text:
        if len(paragraph) == 0:
            output.append("\n\n")
        for word in paragraph:
            if word.lower() in freq_list[:number_of_words]:
                # if the word is in the list, put it unaltered into the output
                output.append(word)
                words_known += 1
            elif word.isdigit():
                output.append(word)
            else:
                # if it's not in the list, censor it to show it isn't a known word
                output.append('x' * len(word))
                words_not_known += 1

    return output

def listToString(list_input,delimitor):
    return (delimitor.join(list_input))

def getNumber():
    valid = False
    result = 0
    while valid is False:
        number_of_words = input("We will show the top x words. \n How many words to be included? \n\n")
        try:
            result = int(number_of_words)
            if result <= 10000 and result >= 0:
                valid = True
            else:
                print("Please choose a number between 0 and 10,000")
        except ValueError:
            print("That's not an integer, please try again")
    print()
    return result

def getLanguage():
    valid = False
    language = ""
    while valid == False:
        response = input("What language are you working in?\t")
        if response.lower() == "es" or response == "Spanish":
            language = "es"
            valid = True
            print()
            return language
        elif response.lower() == "en" or response == "English":
            language = "en"
            valid = True
            print()
            return language
        else:
            "Please enter a valid language. \n"

def getLanguageList(getLanguage):
    if getLanguage == "es":
        return freq_list
    elif getLanguage == "en":
        return english_freq_list
    else:
        print("Error getting frequency list.")

def end_summary():
    total_words = words_known + words_not_known
    percentage_known = int((words_known/total_words)*100)
    
    return(f"If you knew the top {freq_num} words, you would be able to understand {words_known} out of {total_words} words in this text. \nThat's {percentage_known}%!\n")

freq_list = open("txts/10000_formas.txt", encoding="utf-8").readlines()
english_freq_list = open("txts/English_10000_words.txt", encoding="utf-8").readlines()
words_known = 0
words_not_known = 0
freq_num = 0

input_text_h = open("txts/input.txt", encoding='utf-8')
input_text = input_text_h.readlines()
doc = docx.Document()

clean_input = input_to_list(input_text)

current_list = getLanguageList(getLanguage())
# asks which langugage to use the frequency list from
current_list = [x.replace('\n', '') for x in current_list]
# removes newlines from each item in list

output = compare_lists(clean_input, current_list, getNumber())
output_string = listToString(output," ")
print(output_string) # prints to console
doc.add_paragraph(output_string) # saves to document

summary = end_summary()

doc.add_paragraph(summary)
print()
print(summary)

try:
    doc.save('output.docx')
except PermissionError:
    print("Document likely still open in word")